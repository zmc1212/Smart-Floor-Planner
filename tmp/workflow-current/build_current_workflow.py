from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(r"G:\workspace\向总\Smart-Floor-Planner")
REFERENCE = Path(r"G:\workspace\向总\家客来AI获客系统完整工作流.docx")
OUTPUT = ROOT / "output" / "家客来AI获客系统-当前能力工作流-V1.0.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if i == 1 and str(value) == "Implemented":
                set_cell_shading(cells[i], "E2F0D9")
            elif i == 1 and str(value) == "Limited":
                set_cell_shading(cells[i], "FFF2CC")
            elif i == 1 and str(value) == "Placeholder":
                set_cell_shading(cells[i], "FCE4D6")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Pt(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_styles(doc):
    styles = doc.styles
    for name, size, color, bold in [
        ("Normal", 10.5, None, False),
        ("Heading 1", 17, "1F4E79", True),
        ("Heading 2", 13.5, "2E74B5", True),
        ("Heading 3", 12, "333333", True),
    ]:
        s = styles[name]
        s.font.name = "微软雅黑"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        s.font.size = Pt(size)
        s.font.bold = bold
        if color:
            s.font.color.rgb = RGBColor.from_string(color)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.15


def build():
    doc = Document(str(REFERENCE))
    clear_body(doc)
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Pt(59)
    section.bottom_margin = Pt(59)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("《家客来AI获客系统》当前能力工作流")
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("基于现有代码、数据库与模块契约的运行基线")
    add_body(doc, "文档版本：V1.0（当前实现基线）")
    add_body(doc, "适用范围：管理端 + 原生微信小程序；不将规划功能、演示文案或历史路线图视为已实现能力。")
    add_body(doc, "使用方式：本文件可作为当前产品培训、测试用例编写和后续需求评审的工作流基线；新增能力必须先更新状态、数据契约和权限边界。")

    add_heading(doc, "文档概要", 1)
    add_table(doc, ["项目", "当前基线"], [
        ("产品形态", "Next.js 管理端 + 原生微信小程序 + PostgreSQL 租户数据 + BLE 测量设备"),
        ("当前核心闭环", "线索创建 → 设计师/测量员协作 → 正式量房 → AI 设计工作流 → 跟进/签约状态记录"),
        ("已实现获客协作", "测量员绑定设计师、线索自动归属、设计师获客确认、固定金额获客提成、站内/微信通知"),
        ("当前不包含", "外部推荐人裂变、定点活动经营、抖音/小红书等平台全量归因、个人微信聊天读取、自动代发佣金"),
    ], [130, 360])

    add_heading(doc, "第一部分：参考工作流与当前实现对照", 1)
    add_body(doc, "本部分以《家客来AI获客系统完整工作流》V7.0 的步骤、渠道和机制为参照，逐项对照当前代码。状态只表示当前程序是否有真实页面、API、数据和权限路径，不表示商业规则已经被业务方最终确认。")
    add_heading(doc, "1.1 统一主流程对照", 2)
    add_table(doc, ["V7.0 目标步骤", "当前程序对应能力", "状态", "差距与是否需要实现"], [
        ("0 招募认证", "企业员工由管理员创建/维护；测量员与设计师绑定", "Limited", "没有推荐人实名认证、独立第三方协议、培训考核、外部合作方资格链路；若保留推荐人渠道，需要新增。"),
        ("1 触达推荐", "员工/测量员在小程序创建业主线索；线索表保存 source", "Limited", "没有推荐人专属码、业主自助授权留资和多渠道首触达事件；需要统一归因中台。"),
        ("2 提交线索", "POST /api/leads，手机号去重，服务端写入 promoter/assignedTo/new", "Implemented", "当前满足企业内部线索录入；不等同于 V7.0 推荐人提交和跨渠道冲突裁定。"),
        ("3 AI 分级", "线索状态、跟进记录、AI 设计能力存在", "Placeholder", "没有 S/A/B/C 可解释评分、问卷、B 级孵化池 Owner 和触达 SOP；需要新增。"),
        ("4 派单测绘", "测量员绑定设计师；正式 surveying-editor 可量房", "Limited", "没有设计师自主派单、预约、接单/改约/爽约、测绘任务履约证据和共享测绘池；需要新增。"),
        ("5 AI 跟进", "线索详情跟进记录 + AI 设计工作流", "Limited", "设计生成已实现；个人微信话术、情绪、话单质检没有数据源，必须改为 CRM/企微授权增强。"),
        ("6 转化签约", "线索 converted 状态、AI 产物和跟进记录", "Limited", "没有完整在线签约、L3 方案深化图签发、客户确认和满意度闭环；需要明确边界后新增。"),
        ("7 结算分配", "获客提成事务生成、后台人工 paid；订单提成另有流程", "Limited", "固定金额获客提成已实现；双段式推荐人佣金、自动支付、税务与分账未实现。"),
        ("8 考核复盘", "基础工作台、线索/户型/任务汇总、部分 SLA 提醒", "Limited", "没有 V7.0 定义的渠道 ROI、红黄绿灯、赛马、单位经济和人才盘点；需要真实成本/收入事件后实现。"),
    ], [90, 170, 70, 210])
    add_heading(doc, "1.2 三大渠道对照", 2)
    add_table(doc, ["参考渠道", "当前可复用能力", "当前缺失", "建议"], [
        ("编外线下网络", "线索创建、设计师归属、获客确认、提成通知", "推荐人身份、推广码、首触达归因、协议/税务、双段佣金、防刷", "暂不按完整渠道上线；先做合作方和归因基础模型"),
        ("自营定点活动", "线索表、企业报备、跟进记录、通知", "活动、场次、物料码、到场/加微/预约/测绘/签约漏斗、活动成本", "若试点需要，先做最小活动登记和来源码，不先做完整 ROI"),
        ("线上平台", "小程序留资、AI/量房后续流程", "抖音/小红书/美团等平台订单、UTM、体验券、核销、退款、客服 SLA", "先统一到短链/二维码 + 小程序授权留资，平台 API 后置"),
    ], [110, 150, 180, 100])
    add_heading(doc, "1.3 四层管理对照", 2)
    add_table(doc, ["参考层", "当前实现", "结论"], [
        ("AI 辅助", "AI 生成、工作流阶段、额度和结果媒体", "设计生产辅助已实现；销售会话 AI 未实现"),
        ("设计师主观能动性", "设计师负责线索确认、跟进和 AI 方案操作", "部分实现；没有小组经营、排名、奖励再分配"),
        ("系统规则引擎", "租户/RLS、状态、归属快照、通知去重、提成幂等、归档保护", "基础规则已实现；没有 V7.0 归因、风控、孵化和成本规则"),
        ("老板总控", "企业管理、员工、规则配置、提成结算、基础工作台", "基础管理已实现；完整经营驾驶舱和利润决策未实现"),
    ], [120, 240, 180])
    add_heading(doc, "1.4 参考机制缺口与优先级", 2)
    add_table(doc, ["V7.0 机制", "当前状态", "优先级", "需要实现的最小内容"], [
        ("三个免费钩子", "Partial", "P1", "明确免费服务边界；把 AI 产物标记为设计沟通成果，不承诺施工图"),
        ("专属推广码与渠道 ROI", "Placeholder", "P1", "渠道/活动/推荐人码、不可变归因事件、冲突仲裁、成本事件"),
        ("第一触达锁定", "Partial", "P1", "改称系统归因冻结；保存扫码/授权/提交事件并提供申诉"),
        ("双段式激励", "Placeholder", "P1", "佣金规则版本、触发条件、冻结期、作废、税务结算主体和审计"),
        ("AI 分级与孵化", "Placeholder", "P2", "结构化问卷、可解释规则评分、Owner、T+1/T+3/T+7 任务和升级路径"),
        ("AI 验真", "Partial", "P2", "先做量房完整性/设备/时间异常提示和人工复核，不自动认定作弊"),
        ("三级效果图", "Partial", "P1", "L1/L2/L3 改成概念/精细/深化方案，增加水印、客户确认和人工审核"),
        ("红黄绿灯/赛马", "Placeholder", "P3", "先定义 cohort、成本、收入、满意度和权重，再做看板"),
        ("双向评价", "Placeholder", "P2", "客户评价、设计师评价合作方、匿名/权限和反报复规则"),
        ("合规底线", "Partial", "P1", "单独同意、撤回、删除/匿名化、协议版本和操作审计"),
    ], [130, 80, 55, 275])

    add_heading(doc, "第二部分：核心模式总览", 1)
    add_heading(doc, "1.1 产品核心理念", 2)
    add_body(doc, "当前系统的真实定位：以客户线索为业务主线，把设计师、测量员、正式 version-4 户型数据和 AI 设计产物放在同一企业租户边界内，支持从线索创建到方案设计的可追踪协作。")
    add_body(doc, "系统中的 AI 当前主要承担设计生成、方案阶段推进、素材管理和额度控制；它不是个人微信聊天机器人，也不自动替代设计师判断。")
    add_heading(doc, "1.2 当前状态定义", 2)
    add_table(doc, ["状态", "含义", "当前处理"], [
        ("Implemented", "页面、API、数据库和权限边界均有真实运行路径", "可作为当前功能承诺"),
        ("Limited", "真实路径存在，但依赖登录、企业配置、BLE、外部 AI/通知服务或人工操作", "可试用，需标注前置条件"),
        ("Placeholder", "仅有入口、文案或演示结构，没有真实后端闭环", "不得作为已交付能力宣传"),
    ], [110, 190, 190])
    add_heading(doc, "1.3 当前统一流程图", 2)
    add_body(doc, "登录/企业上下文 → 测量员或授权员工创建线索 → 服务端按绑定关系写入设计师负责人 → 站内通知设计师 → 设计师确认获客（独立审计事实） → 线索进入量房/设计阶段 → BLE 或手动完成正式 version-4 测量 → AI 工作流生成和选择方案产物 → 跟进记录/签约状态 → 归档与复盘。")

    add_heading(doc, "第二部分：当前实际执行流程", 1)
    add_heading(doc, "2.1 线索与获客协作流程", 2)
    add_table(doc, ["步骤", "事项", "执行人", "当前规则"], [
        ("0", "登录与上下文恢复", "员工/管理员", "小程序通过 /api/auth/miniprogram 恢复 JWT；所有企业数据按租户和角色隔离。"),
        ("1", "创建客户线索", "测量员或授权员工", "提交姓名、手机号、小区等信息；服务端忽略客户端负责人和状态，使用当前登录人及当前绑定关系。"),
        ("2", "自动归属与通知", "系统", "新线索写入 new，并按 measurer_designer_bindings 写入设计师负责人；通知先写站内，再尝试微信订阅消息。"),
        ("3", "获客确认", "负责设计师", "设计师在 acquisition-center 确认已完成微信交接；只写 acquired_at/acquired_by，不改变线索主状态。"),
        ("4", "获客提成生成", "系统/管理员", "确认与唯一待结算获客提成在同一事务中完成；金额为企业固定配置快照。"),
        ("5", "线索跟进", "设计师/测量员", "使用线索详情、跟进记录、通知和任务工作台；不读取个人微信聊天内容。"),
    ], [45, 110, 100, 235])
    add_heading(doc, "2.2 正式量房流程", 2)
    add_body(doc, "正式测量的唯一入口是 surveying-editor。客户线索与正式户型通过 leadId/floorPlanId 关联，量房数据使用 version: 4、measurementMode: surveying、surveyGraph 的结构；BLE 读数写入测量审计，未完成首次云端保存时保持排队。")
    add_table(doc, ["阶段", "当前能力", "状态"], [
        ("准备", "设备绑定、授权、最近户型继续、线索关联", "Implemented / Limited"),
        ("采集", "BLE 激光测距、手动输入、墙图 Canvas、房间和墙体编辑", "Implemented；依赖设备时 Limited"),
        ("保存", "正式 floor plan 保存、测量记录和审计写入 PostgreSQL", "Implemented"),
        ("校验", "数据结构与墙图规则校验；不是 GPS 轨迹或 AI 反作弊判定", "Implemented / 规则校验"),
        ("交付", "将正式户型作为 AI 设计输入或后台只读查看", "Implemented"),
    ], [100, 300, 90])
    add_heading(doc, "2.3 AI 设计流程", 2)
    add_body(doc, "AI 设计由正式户型、图片或历史产物作为输入，创建工作流并按阶段执行。执行前需要明确确认，系统记录 provider attempt、媒体、任务状态和额度冻结/消耗/释放。生成结果是设计沟通产物，不等同于施工图纸。")
    add_table(doc, ["能力", "当前状态", "边界"], [
        ("户型/空间生成", "Implemented", "基于正式户型或来源图生成概念设计产物，依赖已配置 AI provider。"),
        ("风格与软装工作流", "Implemented", "支持工作流阶段、基准产物、重试和历史；provider/storage 异常会进入失败或待处理状态。"),
        ("AI 额度与计费", "Implemented", "企业额度、价格快照、幂等结算和失败释放已实现。"),
        ("聊天话术/情绪/质检", "Placeholder", "当前没有个人微信聊天读取链路，不应作为当前功能宣传。"),
        ("施工级效果图", "Placeholder", "当前没有施工图签发、材料清单和持证人员审核闭环。"),
    ], [145, 90, 255])

    add_heading(doc, "第三部分：角色与权限边界", 1)
    add_table(doc, ["角色", "当前可做", "当前不可推定"], [
        ("企业管理员", "员工、绑定、租户规则、线索、量房、AI 配置和提成结算管理", "不能据此推定已具备全渠道 ROI 或自动发薪"),
        ("设计师", "查看负责人线索、确认获客、跟进、选择测量结果、发起 AI 设计", "不能读取私人微信聊天；不能任意修改历史线索负责人"),
        ("测量员", "创建线索、查看自己的协作任务、查看绑定设计师、完成正式量房", "不是文档 V7.0 中可自由调度的共享测绘池"),
        ("渠道地推 salesperson", "使用企业报备/推广记录流程，处理企业合作线索", "该角色不是当前业主推荐人 N1，不能将企业报备记录当作业主获客线索"),
        ("业主/客户", "可通过留资、查看设计成果等有限路径参与", "当前没有完整的客户进度中心、在线签约和满意度闭环"),
    ], [110, 220, 160])
    add_body(doc, "权限原则：API 使用 shared tenant helpers、Mini Program staff context 和 PostgreSQL RLS；客户端传入的负责人、状态或企业 ID 不作为可信来源。")

    add_heading(doc, "第四部分：模块能力详解", 1)
    add_heading(doc, "4.1 线索管理与归档", 2)
    add_body(doc, "管理端 /leads 与小程序 leads-management 提供线索列表、搜索、四步业务状态、跟进记录、正式户型关联、详情和归档生命周期。主状态为 new → measuring → designing → converted，closed 为终止状态；获客确认是独立事实，不新增 acquired 主状态。")
    add_heading(doc, "4.2 测量员—设计师协作", 2)
    add_body(doc, "当前闭环解决的是企业内部测量员与设计师的交接，不是外部推荐人裂变。一个测量员只有一个当前绑定设计师；换绑只影响后续新线索，历史线索保留创建时的负责人快照。")
    add_heading(doc, "4.3 提成与通知", 2)
    add_body(doc, "获客提成独立于订单提成，确认时快照企业固定金额，初始为 pending_settlement；当前后台可人工标记 paid，不含支付渠道或银行代发。通知按去重键落库，微信失败不回滚业务事务。")
    add_heading(doc, "4.4 企业报备流程的范围", 2)
    add_body(doc, "promotion-records 是企业合作报备、公共池、保护期、认领、跟进、分配和提醒流程。它与业主线索获客是两套业务对象，当前不能直接支撑 V7.0 所称的推荐人专属推广码、业主渠道归因或活动 ROI。")

    add_heading(doc, "第五部分：当前已落地的关键机制", 1)
    add_table(doc, ["机制", "当前实现", "限制"], [
        ("租户与角色隔离", "PostgreSQL RLS、企业上下文和行级访问", "需要有效登录和企业配置"),
        ("固定归属快照", "创建时写入 promoter/assigned_to，禁止客户端覆盖", "不是跨渠道现实世界首触达证明"),
        ("手机号去重", "命中已有线索返回已有数据，不重复通知/提成", "不是完整设备、微信、楼栋复合反作弊"),
        ("获客确认幂等", "线索条件更新与唯一提成事务化，支持并发保护", "金额规则当前是企业固定金额"),
        ("归档保护", "归档保留户型、AI、提成、通知和跟进；写入返回 LEAD_ARCHIVED", "永久删除有严格保护关系限制"),
        ("AI 额度生命周期", "冻结、消耗、释放、provider attempt、媒体结果和重试", "依赖外部 provider/storage"),
        ("通知可靠性", "先站内写入，再记录微信 sent/failed/skipped", "不是预约成功或客户已读证明"),
    ], [120, 235, 135])

    add_heading(doc, "第六部分：当前功能矩阵", 1)
    add_table(doc, ["领域", "能力", "状态", "当前说明"], [
        ("线索", "创建、列表、详情、跟进、状态、归档", "Implemented", "管理端和小程序均有真实 API/页面"),
        ("协作", "测量员绑定设计师、获客确认、协作任务", "Implemented", "限企业内部绑定关系"),
        ("提成", "获客提成生成、查询、人工结算", "Implemented / Limited", "自动支付和银行代发未实现"),
        ("通知", "站内通知、订阅消息、去重、日志", "Implemented / Limited", "微信模板和授权是外部前置条件"),
        ("正式量房", "BLE、手动测量、墙图、version-4 保存", "Implemented / Limited", "BLE 硬件和设备绑定是前置条件"),
        ("AI 设计", "工作流、阶段、生成、结果、历史、额度", "Implemented / Limited", "依赖 provider、媒体存储和企业额度"),
        ("企业报备", "公共池、认领、保护期、跟进和提醒", "Implemented", "对象是企业合作报备，不是业主推荐人"),
        ("推荐人裂变", "专属码、推荐人收益、首触达归因", "Placeholder", "当前没有对应业主渠道数据模型"),
        ("活动管理", "活动策划、到场、ROI", "Placeholder", "当前没有活动对象和漏斗统计"),
        ("平台订单", "0.01 元体验券、核销、退款", "Placeholder", "当前没有多平台统一订单闭环"),
        ("私聊 AI", "话术、情绪、话单质检、经验萃取", "Placeholder", "没有个人微信聊天数据源"),
        ("老板驾驶舱", "渠道 ROI、红黄绿灯、赛马、利润", "Limited", "现有首页/工作台是基础汇总，不是 V7.0 完整驾驶舱"),
    ], [85, 145, 100, 160])

    add_heading(doc, "第七部分：当前可用指标", 1)
    add_body(doc, "当前可以可靠统计的指标必须来自已有事实表和事件：线索数、按状态线索数、正式户型数、测量记录数、获客确认数、待结算/已结算获客提成、AI 任务状态和通知状态。")
    add_table(doc, ["指标", "当前口径", "状态"], [
        ("线索状态漏斗", "按 leads.status 统计 new/measuring/designing/converted/closed", "Implemented"),
        ("获客确认率", "有 acquired_at 的线索 / 线索 cohort", "Implemented；需明确时间窗口"),
        ("测绘数据量", "formal floor plans 与 measurements 记录", "Implemented"),
        ("AI 生成成功率", "generation 状态与 provider attempt 结果", "Limited；受 provider 影响"),
        ("渠道 ROI", "需要渠道成本、签约收入和统一归因", "Placeholder"),
        ("客户满意度", "当前没有完整客户匿名评价事实表", "Placeholder"),
        ("小组排名", "当前没有 V7.0 定义的综合得分和公开赛马", "Placeholder"),
    ], [120, 300, 90])

    add_heading(doc, "第八部分：当前边界与后续建议", 1)
    add_heading(doc, "8.1 当前不能从系统推导的能力", 2)
    add_bullet(doc, "不能从现有字段推导外部推荐人、活动、平台投流的真实首触达和 ROI。")
    add_bullet(doc, "不能读取个人微信私聊，因此不能承诺自动情绪感知、话单质检或成功经验萃取。")
    add_bullet(doc, "不能把 BLE 测量记录当作 GPS 轨迹验真，也不能仅凭设备号自动认定作弊。")
    add_bullet(doc, "不能把 AI 概念效果图表述为施工图纸或施工级交付。")
    add_bullet(doc, "不能把人工标记 paid 解释为自动支付或合规代发。")
    add_heading(doc, "8.2 如果扩展 V7.0，建议先补的基础对象", 2)
    add_body(doc, "只有在业务、法务和财务规则定稿后，才建议新增：合作方/推荐人档案、不可变归因事件、客户授权记录、预约与测绘任务、活动与平台订单、佣金规则版本、争议仲裁记录、客户评价和成本事件。")
    add_body(doc, "这些对象应与现有 leads、formal floor plans、AI generations 和 acquisition commissions 通过明确关系连接，不能继续复用 source、promoterId 或 follow_up_records 这类已有字段承载多个概念。")

    add_heading(doc, "第九部分：版本与依据", 1)
    add_table(doc, ["依据", "用途"], [
        ("docs/admin-system-modules.md", "管理端当前页面、API、权限和状态基线"),
        ("docs/miniprogram-system-modules.md", "小程序页面、角色、正式量房和 AI 入口基线"),
        ("docs/measurer-designer-acquisition.md / .zh-CN.md", "测量员—设计师获客确认、提成和通知契约"),
        ("docs/surveying-module/formal-surveying.md", "正式 version-4 wall graph、BLE 审计和保存边界"),
        ("admin/src/db/schema.ts", "PostgreSQL leads、floor_plans、measurements、commission、notification 等表结构"),
    ], [240, 270])
    add_body(doc, "文档状态：当前运行基线；不等同于 V7.0 目标蓝图。任何功能新增、路由变更、权限变更或数据契约变更，都必须同步更新本文件与对应中英文模块清单。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build()
