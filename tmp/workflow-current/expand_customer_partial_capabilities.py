from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(r'G:\workspace\向总\Smart-Floor-Planner\output\家客来AI获客系统-客户版工作流对照-V1.1.docx')
TARGETS = [
    Path(r'G:\workspace\向总\Smart-Floor-Planner\output\家客来AI获客系统-客户版工作流对照-V1.1.docx'),
    Path(r'G:\workspace\向总\Smart-Floor-Planner\output\家客来AI获客系统-客户版工作流对照-V1.0.docx'),
]


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    element = props.find(qn('w:shd'))
    if element is None:
        element = OxmlElement('w:shd')
        props.append(element)
    element.set(qn('w:fill'), fill)


def set_cell(cell, text, bold=False, color=None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_after(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, True, 'FFFFFF')
        shade(table.rows[0].cells[index], '1F4E79')
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], value)
            if index == 1:
                shade(cells[index], 'E2F0D9')
            if index == 2:
                shade(cells[index], 'FFF2CC')
    doc.add_paragraph()


doc = Document(str(SOURCE))

# Expand the customer-facing main flow rows, rather than leaving the word
# "partial" to carry the explanation by itself.
main = doc.tables[1]
updates = {
    '0. 招募与认证': (
        '企业可维护设计师和测量员资料，并指定测量员协作的设计师；人员变更后，新客户会按新的协作关系进入流程。',
        '推荐人招募、实名认证、合作协议、培训考核、暂停合作和收益资格，适合作为“合作推荐人管理”单独建设。',
    ),
    '1. 触达推荐': (
        '工作人员可以在移动端录入客户资料，客户进入后续的设计师承接、量房和方案流程。',
        '推荐人二维码、活动二维码、线上投放链接和客户来源追踪，需要在渠道推广前统一建设。',
    ),
    '4. 派单测绘': (
        '测量员可以完成正式量房，量房结果可直接供设计师用于后续方案设计；设计师和测量员之间已有固定协作关系。',
        '预约时间、测量员接单、改约、爽约处理、区域共享测量员池和空跑规则，需要在规模调度前补齐。',
    ),
    '5. AI 跟进': (
        '设计师可查看客户资料和跟进记录，并基于户型或现场图片生成 AI 设计方案，用于客户沟通。',
        '自动客户分级、提醒节奏、话术建议和沟通质量分析，需要基于客户授权的沟通记录逐步建设。',
    ),
    '6. 转化签约': (
        '客户可按“新线索、量房中、方案设计、已签约”持续推进，方案成果和沟通记录能够沉淀。',
        '线上签约、客户确认、满意度评价、深化设计交付和施工交接，需要结合企业交付标准继续建设。',
    ),
    '7. 结算分配': (
        '设计师确认客户交接后，系统可生成测量员的待结算获客提成；企业管理人员可查看并确认发放。',
        '推荐人即时奖励与成交奖励、自动支付、税务处理、退款作废和争议处理，需要建立独立的财务与合规流程。',
    ),
    '8. 考核复盘': (
        '企业可查看客户线索、量房、方案任务和部分业务提醒，掌握团队当前的基本进展。',
        '渠道投入产出、签约利润、团队排名、预警灯和人才盘点，需要在产生稳定的渠道成本与签约数据后建设。',
    ),
}
for row in main.rows[1:]:
    key = row.cells[0].text.strip()
    if key in updates:
        set_cell(row.cells[1], updates[key][0])
        set_cell(row.cells[3], updates[key][1])

# Add a compact detailed section after the existing document so a decision-maker
# can scan exactly what "partial" means without interpreting technical labels.
heading = doc.add_paragraph(style='Heading 2')
heading.add_run('九、“部分具备”能力的具体范围')
lead = doc.add_paragraph(style='Normal')
lead.add_run('以下说明用于帮助项目负责人快速判断首期可以直接采用的范围，以及进入下一阶段前需要确认的业务规则。')
add_table_after(doc, ['业务环节', '现阶段已经具备', '进入下一阶段需补齐'], [
    ('人员协作', '设计师和测量员之间可建立稳定协作关系；客户线索会进入对应设计师的工作范围。', '外部推荐人合作、资格管理、协议签署、培训与收益管理。'),
    ('客户承接', '可录入客户资料、识别重复信息、由设计师承接并持续记录进展。', '客户自主扫码留资、来源确认、跨渠道归属和争议处理。'),
    ('智能量房', '可进行正式量房，测量数据可直接进入户型和设计流程。', '预约排期、多人调度、改约爽约、区域拼单和空跑规则。'),
    ('AI 方案沟通', '可基于户型或图片生成方案，用于风格沟通、空间展示和方案迭代。', '客户分级、自动培育、授权沟通分析、深化方案审核和客户确认。'),
    ('获客收益', '可形成测量员获客提成记录，并由企业进行人工结算确认。', '推荐人双段收益、自动付款、税务处理、退款作废和仲裁机制。'),
    ('经营管理', '可查看客户、量房和方案的基础进展，帮助管理人员掌握日常协作。', '活动与平台数据、渠道成本、签约收入、利润分析、排名和预警。'),
])

for target in TARGETS:
    doc.save(str(target))
print('\n'.join(str(path) for path in TARGETS))
