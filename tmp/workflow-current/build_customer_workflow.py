from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(r'G:\workspace\向总\Smart-Floor-Planner')
REFERENCE = Path(r'G:\workspace\向总\家客来AI获客系统完整工作流.docx')
OUTPUT = ROOT / 'output' / '家客来AI获客系统-客户版工作流对照-V1.0.docx'

def shade(cell, fill):
    p = cell._tc.get_or_add_tcPr()
    el = p.find(qn('w:shd'))
    if el is None:
        el = OxmlElement('w:shd')
        p.append(el)
    el.set(qn('w:fill'), fill)

def cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(text))
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(doc, headers, rows, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, True)
        shade(t.rows[0].cells[i], '1F4E79')
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
            if status_col is not None and i == status_col:
                shade(cells[i], {'已支持':'E2F0D9','部分支持':'FFF2CC','尚未支持':'FCE4D6','建议后续建设':'DDEBF7'}.get(str(value), 'FFFFFF'))
    doc.add_paragraph()
    return t

def heading(doc, text, level):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.add_run(text)

def body(doc, text):
    p = doc.add_paragraph(style='Normal')
    p.add_run(text)

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

def build():
    doc = Document(str(REFERENCE))
    b = doc._element.body
    sect = b.sectPr
    for child in list(b):
        if child is not sect:
            b.remove(child)
    for name, size, color, bold in [('Normal',10.5,None,False),('Heading 1',17,'1F4E79',True),('Heading 2',13.5,'2E74B5',True),('Heading 3',12,'333333',True)]:
        s = doc.styles[name]
        s.font.name = '微软雅黑'
        s._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        s.font.size = Pt(size)
        s.font.bold = bold
        if color: s.font.color.rgb = RGBColor.from_string(color)
    doc.styles['Normal'].paragraph_format.space_after = Pt(5)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.15
    title = doc.add_paragraph(style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('《家客来AI获客系统》客户版工作流对照说明')
    sub = doc.add_paragraph(style='Subtitle')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('参考 V7.0 目标流程与当前系统实际能力的业务化对照')
    body(doc, '版本：V1.0（客户沟通版）')
    body(doc, '阅读说明：本文件不描述代码和技术细节，只回答三个问题：参考方案想实现什么、当前系统已经能做什么、还需要建设什么。')
    body(doc, '状态说明：已支持 = 当前可以真实使用；部分支持 = 主流程能用，但需要人工参与或有前置条件；尚未支持 = 当前不能对客户承诺。')

    heading(doc, '一、先说结论', 1)
    body(doc, '当前系统已经能够支持一条稳定的“客户线索—设计师协作—正式量房—AI设计方案”工作流，适合作为装企内部的客资和设计协作工具。')
    body(doc, '参考 V7.0 方案还增加了推荐人裂变、定点活动、线上平台、自动分级、自动跟进、佣金分账和经营驾驶舱等经营体系。这些内容目前不能全部按成品功能对外承诺，需要分阶段建设。')
    table(doc, ['参考方案的完整目标', '当前系统现实情况', '判断'], [
        ('从多个渠道持续获客并自动分配', '已有线索录入、设计师负责和测量协作', '部分支持'),
        ('用 AI 做获客、跟进和成交辅助', '已有 AI 设计生成和方案工作流', '部分支持'),
        ('老板看渠道投入、转化和利润', '已有基础工作台和业务汇总', '部分支持'),
        ('推荐人、活动、平台和佣金完整闭环', '相关业务对象和规则尚未完整建立', '尚未支持'),
    ], 2)

    heading(doc, '二、主流程对照', 1)
    body(doc, '下面按照参考文档的 0-8 步骤逐项说明。')
    table(doc, ['参考方案步骤', '当前系统能做什么', '状态', '还缺什么'], [
        ('0. 招募与认证', '可以由企业管理人员维护设计师、测量员，并建立两者的协作关系。', '部分支持', '推荐人注册、实名认证、协议签署、培训考核还没有形成完整流程。'),
        ('1. 触达推荐', '工作人员可以在小程序中录入客户线索。', '部分支持', '没有推荐人专属二维码、活动码和多平台来源追踪。'),
        ('2. 提交线索', '可以记录客户姓名、电话、小区等信息，并对重复电话进行处理。', '已支持', '当前主要适用于企业内部人员录入，不等于完整的外部推荐渠道。'),
        ('3. AI 分级', '可以查看线索状态、跟进记录，并使用 AI 设计能力。', '尚未支持', '没有成熟的 S/A/B/C 客户分级、自动问卷和 B 级客户培育机制。'),
        ('4. 派单测绘', '测量员可以按照已有协作关系完成正式量房，并将户型交给设计使用。', '部分支持', '没有完整的预约、接单、改约、爽约处理和共享测量员调度。'),
        ('5. AI 跟进', '可以记录客户跟进，并根据正式户型或图片生成设计方案。', '部分支持', '不能自动读取设计师私人微信聊天，也不能承诺自动情绪分析和话术质检。'),
        ('6. 转化签约', '可以把客户推进到“方案设计”和“已签约”等业务阶段，保存设计成果。', '部分支持', '在线签约、客户确认、满意度评价和正式施工交付还需要补充。'),
        ('7. 结算分配', '设计师确认获客后，可以产生待结算的固定金额获客提成，后台人工确认发放。', '部分支持', '推荐人双段佣金、自动支付、税务处理和争议仲裁尚未建立。'),
        ('8. 考核复盘', '可以查看线索、量房、设计任务和部分提醒信息。', '部分支持', '渠道成本、签约收入、利润、赛马排名和人才盘点还不完整。'),
    ], 2)

    heading(doc, '三、三大获客渠道对照', 1)
    table(doc, ['参考渠道', '当前可利用的能力', '当前缺失', '建议'], [
        ('编外线下网络', '已有客户线索、设计师承接、测量和获客提成基础。', '推荐人身份、专属码、首触达确认、协议、税务和防刷。', '先建设“合作推荐人基础版”，再谈裂变和长期佣金。'),
        ('自营定点活动', '已有客户录入、跟进和通知，可作为活动后的客户承接工具。', '活动计划、场次、现场登记、活动码、投入产出统计。', '试点时先做活动登记和来源标记，不先承诺完整 ROI。'),
        ('线上平台', '客户可以通过小程序进入留资、量房和设计流程。', '各平台订单、体验券、核销、退款和统一客服承接。', '先用统一二维码/短链接引导客户进入小程序，平台深度接入后置。'),
    ])

    heading(doc, '四、当前系统已经可以交付的完整体验', 1)
    table(doc, ['客户或员工角色', '可以完成的事情', '客户能感受到的价值'], [
        ('设计师', '接收负责的线索、确认客户交接、查看客户资料、记录跟进、发起 AI 设计。', '客户资料集中，不用在多个工具之间重复查找。'),
        ('测量员', '录入客户、查看绑定设计师、完成正式量房、上传真实测量结果。', '量房数据可以直接进入后续设计，不需要重复录入。'),
        ('企业负责人', '维护员工关系、查看线索和业务进展、设置获客提成、人工确认结算。', '能看到团队当前的客资和任务状态。'),
        ('客户', '留下基本需求，参与量房和方案沟通，查看设计成果。', '从客户信息到户型和方案形成连续记录。'),
    ])
    body(doc, '这条内部协作流程是当前最适合先落地和试点的版本。它不依赖推荐人裂变，也不依赖读取私人微信聊天，能够独立运行。')

    heading(doc, '五、参考方案中还需要建设的部分', 1)
    table(doc, ['建设内容', '要解决的问题', '优先级', '适合何时建设'], [
        ('推荐人合作管理', '谁可以推荐、推荐关系如何确认、收益如何结算。', '高', '正式做外部渠道前'),
        ('统一来源追踪', '客户来自哪个人、哪个活动、哪个平台，出现冲突时如何裁定。', '高', '所有外部渠道前'),
        ('客户授权与隐私管理', '客户是否同意接收营销、保存资料和使用 AI 辅助。', '高', '外部渠道和 AI 增强前'),
        ('预约与测量任务', '谁去量房、何时上门、客户爽约怎么办、测量员如何结算。', '高', '共享测量员或规模化试点前'),
        ('客户分级与培育', '高意向客户优先处理，普通客户也不会无人跟进。', '中', '已有稳定线索量后'),
        ('效果图分级交付', '区分快速沟通图、精细方案图和人工深化成果，避免过度承诺。', '高', '对外推广免费效果图前'),
        ('客户评价与满意度', '知道服务质量，形成服务改进依据。', '中', '正式扩大客户量后'),
        ('经营看板与 ROI', '判断哪个渠道值得投入、哪个团队需要调整。', '中', '产生真实成本和签约数据后'),
        ('自动支付与财务分账', '佣金如何付款、退款、作废和处理税务。', '高', '正式开放外部佣金前'),
    ], 2)

    heading(doc, '六、建议的分阶段方案', 1)
    table(doc, ['阶段', '建议范围', '目标'], [
        ('第一阶段：内部协作版', '线索管理、设计师/测量员协作、正式量房、AI 设计、跟进记录、人工提成结算。', '先让装企内部流程稳定跑通。'),
        ('第二阶段：合作推荐版', '推荐人档案、推荐二维码、客户授权、来源确认、推荐收益规则和人工审核。', '验证外部推荐是否真的带来有效客户。'),
        ('第三阶段：规模运营版', '预约调度、活动管理、线上平台承接、客户分级、培育任务和评价体系。', '支持多个小区和多个渠道同时运行。'),
        ('第四阶段：经营决策版', '渠道成本、签约收入、利润、排名、预警、自动结算和高级 AI 辅助。', '帮助老板决定预算、人员和制度。'),
    ])
    body(doc, '建议不要一开始就同时建设四个阶段。第一阶段已经具备较好的程序基础，也最容易验证客户是否真正使用。')

    heading(doc, '七、对客户的准确产品表述', 1)
    bullet(doc, '当前可以对外介绍：客户线索管理、设计师与测量员协作、正式智能量房、AI 设计方案、方案过程记录和获客提成管理。')
    bullet(doc, '当前应谨慎表述：自动获客、AI 自动跟进、AI 情绪识别、施工级效果图、全渠道 ROI、自动佣金发放。')
    bullet(doc, '推荐使用的说法：AI 辅助设计、快速生成沟通方案、量房数据直接进入设计、获客提成可追踪、经营数据逐步完善。')

    heading(doc, '八、客户决策建议', 1)
    body(doc, '如果目标是尽快开始使用，建议以当前已支持的“内部协作版”作为首期产品，不要把 V7.0 全部目标作为首期交付承诺。')
    body(doc, '如果目标是建设完整的外部获客平台，则需要先确认推荐人合作、佣金税务、客户授权、来源归属和效果图交付边界，再进入第二阶段设计。')
    body(doc, '本文件用于帮助客户区分“现在能用的产品”和“未来可以建设的能力”，避免把业务蓝图误解为已经全部上线。')

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)

if __name__ == '__main__':
    build()
