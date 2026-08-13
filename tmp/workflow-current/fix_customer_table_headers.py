from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor


source = Path(r'G:\workspace\向总\Smart-Floor-Planner\output\家客来AI获客系统-客户版工作流对照-V1.1.docx')
target = Path(r'G:\workspace\向总\Smart-Floor-Planner\output\家客来AI获客系统-客户版工作流对照-V1.2.docx')

doc = Document(str(source))
for table in doc.tables:
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.bold = True

doc.save(str(target))
print(target)
