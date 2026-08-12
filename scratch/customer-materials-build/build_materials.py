from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.colors import HexColor, white
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader

ROOT = Path(r"G:/workspace/向总/Smart-Floor-Planner")
OUT = ROOT / "customer-materials"
ASSETS = ROOT / "design-references"
OUT.mkdir(exist_ok=True)

GREEN = "20A36B"
DARK = "17362D"
MINT = "EAF7F0"
MUTED = "5B6B65"

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def borderless(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)

def set_cell_margin(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def add_run(p, text, size=11, color=DARK, bold=False):
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    return r

def set_spacing(p, before=0, after=6, line=1.2):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line

def header(doc, page_no, title):
    p = doc.add_paragraph()
    set_spacing(p, 0, 2)
    add_run(p, "家客来  |  客户推广产品说明书", 9, GREEN, True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2 = doc.add_paragraph()
    set_spacing(p2, 0, 12)
    add_run(p2, title, 23, DARK, True)
    line = p2._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "8"); bottom.set(qn("w:color"), GREEN)
    pbdr.append(bottom); line.append(pbdr)

def footer(doc, page_no):
    p = doc.add_paragraph()
    set_spacing(p, 10, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f"家客来 · 预约试用材料 · {page_no:02d}", 8, MUTED)

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_spacing(p, 0, 4, 1.18)
        add_run(p, item, 10.5, DARK)

def image(doc, relative, width=2.35):
    path = ASSETS / relative
    if path.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))

def block(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borderless(table)
    c = table.cell(0,0); shade(c, MINT); set_cell_margin(c, 150, 180, 150, 180)
    p = c.paragraphs[0]; set_spacing(p,0,4)
    add_run(p, label, 10, GREEN, True)
    p = c.add_paragraph(); set_spacing(p,0,0)
    add_run(p, body, 10.5, DARK)

def page_break(doc):
    doc.add_page_break()

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
sec.top_margin = Inches(.6); sec.bottom_margin = Inches(.6)
sec.left_margin = Inches(.7); sec.right_margin = Inches(.7)
styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

# 1 Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_spacing(p, 24, 8)
logo = ASSETS / "brand-concepts/jiakuke-logo-v1.png"
if logo.exists(): p.add_run().add_picture(str(logo), width=Inches(1.25))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_spacing(p, 4, 8)
add_run(p,"家客来",31,GREEN,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_spacing(p,0,14)
add_run(p,"装修公司的客户经营与空间服务工作台",22,DARK,True)
block(doc,"产品定位","把线索获取、正式量房、设计协作与经营管理，串成可落地的数字化工作闭环。")
image(doc,"brand-concepts/d-fusion-f3-dual-state-v1.png",2.1)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_spacing(p,10,0)
add_run(p,"适用于装修公司负责人、测量员与设计师团队",10,MUTED)
footer(doc,1)

# 2
page_break(doc); header(doc,2,"一套系统，服务两类关键角色")
table=doc.add_table(rows=1,cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER; borderless(table)
for i,(title,items) in enumerate((("企业负责人",["看见线索、量房、设计与交接的过程","统一员工、客户、设备、权限与经营规则","把客户资产沉淀在企业，而不是散落在个人微信"]),("测量员与设计师",["测量员录入客户并完成正式量房","设计师基于正式户型推进方案与客户交接","双方通过通知、回执和协作任务减少口头断点"]))):
    c=table.cell(0,i); shade(c, MINT if i==0 else "F5F8F6"); set_cell_margin(c,180,180,180,180); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    p=c.paragraphs[0]; set_spacing(p,0,7); add_run(p,title,15,GREEN if i==0 else DARK,True)
    for item in items:
        p=c.add_paragraph(style="List Bullet"); set_spacing(p,0,4); add_run(p,item,10.5,DARK)
block(doc,"共同目标","让客户从进入系统，到量房、方案、协作与跟进，都能留下可查看、可协同的业务记录。")
image(doc,"brand-concepts/d-fusion-f1-solid-friendly-v1.png",1.65); footer(doc,2)

# 3
page_break(doc); header(doc,3,"让装修业务从“靠人记”变成“有过程可看”")
bullets(doc,["线索常散在个人微信、表格和口头交接里，负责人难以追踪客户去向。","现场量房、户型信息和后续设计脱节，重复沟通影响效率。","测量员与设计师协作缺少统一任务与通知入口。","客户进度、员工分工、设备使用和经营规则缺少一个企业级视图。"])
block(doc,"家客来的回答","用同一条客户服务链路连接前端现场工作与后台企业管理，让每个角色只看到自己该做的事。")
image(doc,"all-pages-ip-v1/00-overview-core.png",2.55); footer(doc,3)

# 4
page_break(doc); header(doc,4,"从线索进入，到空间服务交付")
steps=[("01","线索进入","记录客户、社区、面积、风格与跟进状态"),("02","正式量房","建立毫米级墙图，记录门窗、空间与测量审计"),("03","AI设计","基于已完成的正式户型选择整户或单空间方案"),("04","协作交接","设计师确认客户交接，测量员查看通知、回执与提成"),("05","企业管理","后台统一员工、权限、设备、线索、AI额度与结算记录")]
for no,step_title,b in steps:
    flow_table=doc.add_table(rows=1,cols=2); borderless(flow_table); flow_table.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=flow_table.cell(0,0); c.width=Inches(.7); shade(c,GREEN); set_cell_margin(c,110,120,110,100); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(p,no,11,"FFFFFF",True)
    c=flow_table.cell(0,1); shade(c,"F5F8F6"); set_cell_margin(c,110,130,110,130); p=c.paragraphs[0]; set_spacing(p,0,2); add_run(p,step_title,11,DARK,True); add_run(p,"  "+b,10,MUTED)
footer(doc,4)

# 5
page_break(doc); header(doc,5,"老板视角：把过程、团队与客户资产放在一起")
bullets(doc,["过程可视：以线索阶段、正式量房状态和设计任务为基础查看客户服务链路。","协同可控：管理员可管理员工角色、企业权限、设备绑定与固定规则。","客户可沉淀：客户、户型、测量记录、方案任务和协作事实围绕企业数据边界保存。","经营有依据：线索、协作和提成结算记录可作为后续统计与管理基础。"])
block(doc,"对外表达建议","不要承诺未验证的转化率或节省比例。使用“帮助团队建立统一过程与数据基础”这类可被真实产品支持的价值表达。")
image(doc,"all-pages-ip-v1/01-home-v2.png",1.75); footer(doc,5)

# 6
page_break(doc); header(doc,6,"测量员与设计师：把现场动作接进企业流程")
items=[("录入客户","登记姓名、电话、社区、面积、风格等真实信息。"),("启动量房","从客户或项目入口进入唯一正式量房页面。"),("建立户型","支持直墙、斜墙、手动或兼容 BLE 测距，保存草稿与正式户型。"),("推动交接","设计师在协作工作台确认交接；测量员可看通知、回执和提成状态。")]
for i,(t,b) in enumerate(items,1):
    p=doc.add_paragraph(); set_spacing(p,2,4); add_run(p,f"{i}. ",12,GREEN,True); add_run(p,t+"：",12,DARK,True); add_run(p,b,10.5,DARK)
image(doc,"all-pages-ip-v1/02-leads-management.png",1.55); footer(doc,6)

# 7
page_break(doc); header(doc,7,"线索与客户档案：让客户推进有清晰的业务轨迹")
bullets(doc,["支持客户信息录入、线索列表、详情查看、搜索与状态筛选。","客户流程按“新线索 → 量房中 → 设计方案 → 已签约”管理；关闭客户为终止状态。","正式量房可与客户关联，并展示项目与空间等真实读模型信息。","归档客户从日常列表隐藏，保留历史资产并通过授权的管理流程恢复。"])
block(doc,"现场展示重点","展示一个样例客户从录入到关联正式户型的过程；不要将设计师交接确认误说成客户生命周期的必经步骤。")
image(doc,"all-pages-ip-v1/08-lead-detail.png",1.55); footer(doc,7)

# 8
page_break(doc); header(doc,8,"正式量房：为后续设计建立可信的空间底图")
bullets(doc,["唯一正式量房页使用 version-4 墙图；墙体、尺寸、门窗和空间统一使用毫米制读写。","支持手动输入与兼容 BLE 激光测距仪；BLE连接取决于手机授权、设备兼容性与实际连接状态。","支持草稿恢复、云端保存、闭合空间、门窗编辑、测量审计和基础3D开口预览。","后台可基于正式墙图提供2D/3D查看与DXF下载；小程序端完整报告导出不作为当前宣传承诺。"])
block(doc,"正确定位","正式量房的价值是减少空间信息断点，不把它宣传为替代人工复核或所有硬件环境均可用的自动测量。")
image(doc,"all-pages-ip-v1/03-surveying-editor-idle.png",1.55); footer(doc,8)

# 9
page_break(doc); header(doc,9,"AI设计：基于正式户型推进方案表达")
bullets(doc,["企业人员可从正式户型进入AI设计，选择整户或单空间作为目标。","工作流覆盖空间基线、风格方案、软装与方案完善等阶段，并保留任务历史、状态与重试入口。","生成任务依赖企业AI配置、额度、有效正式户型和当前权限。","AI结果用于方案参考与概念表达，不应承诺为施工图、自动报价或自动成交结果。"])
block(doc,"演示建议","优先展示已准备好的成功样例与历史结果；现场生成失败时切换至备用视频，不把外部服务状态解释为产品功能缺失。")
image(doc,"all-pages-ip-v1/04-ai-design-home.png",1.55); footer(doc,9)

# 10
page_break(doc); header(doc,10,"协作、管理与试用：从会场进入真实落地")
bullets(doc,["测量员与设计师的获客协作独立于客户业务状态：设计师确认交接后生成待结算提成记录，并通知测量员。","企业后台提供员工、角色/权限、线索、设备与获客提成规则的管理能力；自动支付或银行代发当前未开放。","数据在企业上下文和角色边界内读取；不同角色只访问职责范围内的客户、任务和经营信息。"])
block(doc,"预约试用","建议现场登记企业名称、联系人、角色、现有获客方式、测量设备情况与希望优先体验的功能。\n\n待替换：企业联系人 / 微信或电话 / 预约二维码 / 试用入口")
footer(doc,10)

docx_path = OUT / "家客来-产品说明书-三角色协同优化版.docx"
doc.save(docx_path)

# Matching PDF, designed for mobile sharing. Separate export is necessary in this runtime where LibreOffice is unavailable.
font_path = Path(r"C:/Windows/Fonts/msyh.ttc")
pdfmetrics.registerFont(TTFont("YaHei", str(font_path), subfontIndex=0))
pdf_path = OUT / "家客来-产品说明书-三角色协同优化版.pdf"
W,H=A4
c=canvas.Canvas(str(pdf_path), pagesize=A4)
c.setTitle("家客来产品说明书")
pages=[
 ("家客来", "装修公司的客户经营与空间服务工作台", ["线索获取、正式量房、设计协作与经营管理", "一条闭环，服务负责人、测量员与设计师", "适用于装修公司负责人、测量与设计协同"], "brand-concepts/d-fusion-f3-dual-state-v1.png"),
 ("一套系统，三类企业角色", "负责人看过程，测量员与设计师推进服务", ["负责人：看见客户服务过程，协同团队，沉淀客户资产", "测量员：录入客户、正式量房、提交户型", "设计师：承接客户交接，基于户型推进方案"], "brand-concepts/d-fusion-f1-solid-friendly-v1.png"),
 ("让装修业务有过程可看", "把分散的工作，接成企业自己的服务链", ["线索不再只散在个人微信、表格与口头交接里", "现场量房、户型与后续设计形成可衔接的工作链路", "企业可统一人员、设备、客户和协作事实"], "all-pages-ip-v1/00-overview-core.png"),
 ("一条真实的客户服务闭环", "从客户进入，到空间服务交付", ["线索进入：记录客户与服务状态", "正式量房：建立毫米级墙图与空间信息", "AI设计、协作交接与企业管理围绕真实客户推进"], "all-pages-ip-v1/01-home-v2.png"),
 ("老板视角", "把过程、团队与客户资产放在一起", ["过程可视：线索、量房、设计任务有清晰状态", "协同可控：员工、权限、设备与规则统一管理", "客户可沉淀：资料和协作事实留在企业数据边界内"], "all-pages-ip-v1/01-home-v2.png"),
 ("测量员与设计师", "把每次上门与方案动作接进企业流程", ["测量员：记录客户信息，从客户入口启动正式量房", "设计师：接收客户交接，基于正式户型推进方案", "双方：查看通知、回执与协作任务，减少口头断点"], "all-pages-ip-v1/02-leads-management.png"),
 ("线索与客户档案", "客户推进，围绕清晰的业务状态", ["客户流程：新线索 → 量房中 → 设计方案 → 已签约", "正式量房可与客户关联，展示真实项目与空间信息", "归档客户保留历史资产，通过授权流程恢复"], "all-pages-ip-v1/08-lead-detail.png"),
 ("正式量房与 BLE", "为后续设计建立可信的空间底图", ["毫米级墙图、门窗、闭合空间、测量审计与草稿恢复", "支持手动输入与兼容 BLE 激光测距", "BLE连接取决于兼容硬件、手机授权与现场连接状态"], "all-pages-ip-v1/03-surveying-editor-idle.png"),
 ("AI设计与协作", "让正式户型进入方案表达", ["选择整户或单空间，保留任务历史、状态与重试入口", "AI结果用于方案参考和概念表达，不等同施工图或自动成交", "交接确认生成待结算记录；自动打款当前未开放"], "all-pages-ip-v1/04-ai-design-home.png"),
 ("预约试用", "用真实客户场景，验证是否适合你的团队", ["现场登记企业名称、联系人、角色、设备与优先体验功能", "安排面向企业的试用演示与产品答疑", "待替换：企业联系人 / 微信或电话 / 预约二维码 / 试用入口"], "brand-concepts/d-fusion-f1-solid-friendly-v1.png"),
]
for n,(title,lead,lines,visual) in enumerate(pages,1):
    c.setFillColor(HexColor("#FFFFFF")); c.rect(0,0,W,H,fill=1,stroke=0)
    c.setFillColor(HexColor("#20A36B")); c.rect(0,H-64,W,64,fill=1,stroke=0)
    c.setFillColor(white); c.setFont("YaHei",11); c.drawString(40,H-39,"家客来  |  产品说明书")
    c.setFillColor(HexColor("#17362D")); c.setFont("YaHei",25); c.drawString(42,H-128,title)
    c.setFillColor(HexColor("#20A36B")); c.setFont("YaHei",15); c.drawString(42,H-162,lead)
    c.setFillColor(HexColor("#EAF7F0")); c.roundRect(40,318,306,184,12,fill=1,stroke=0)
    y=466
    c.setFillColor(HexColor("#17362D")); c.setFont("YaHei",12.5)
    for line in lines:
        c.setFillColor(HexColor("#20A36B")); c.circle(64,y+2,3,fill=1,stroke=0)
        c.setFillColor(HexColor("#17362D")); c.drawString(78,y-3,line)
        y-=52
    visual_path=ASSETS / visual
    if visual_path.exists():
        c.drawImage(ImageReader(str(visual_path)),374,248,150,300,preserveAspectRatio=True,anchor='c',mask='auto')
    c.setFillColor(HexColor("#20A36B")); c.roundRect(42,155,513,2,0,fill=1,stroke=0)
    c.setFillColor(HexColor("#5B6B65")); c.setFont("YaHei",9); c.drawRightString(555,40,f"家客来 · 预约试用材料 · {n:02d}")
    c.showPage()
c.save()
print(docx_path)
print(pdf_path)
